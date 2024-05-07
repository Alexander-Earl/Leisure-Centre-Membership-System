from tkinter import *
from Database import *
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from MainApp import *
import os
import tkinter.messagebox


class Report(Frame):

    def __init__(self, parent, controller):
        Frame.__init__(self)
        # Sets the background colour to the hexadecimal value.
        self.config(bg="#97c7f1")
        self.controller = controller

        # This segment of code implements the 'Generate Member Report' banner onto the window.
        self.member_report_photo = PhotoImage(file="Images/Generate Member Report.png")
        self.member_report_title = Label(self, image=self.member_report_photo, height=80, width=725, borderwidth=0, activebackground="#97c7f1")
        self.member_report_title.pack(anchor=E)
        self.member_report_title.config(bg="#97c7f1")

        # This section implements a back button which allows the staff associate to traverse to the previous window.
        self.back_photo = PhotoImage(file="Images/Back Button.png")
        self.back_button = Button(self, image=self.back_photo, command=lambda: controller.show_frame("Staff"),
                        height=70, width=120, borderwidth=0, activebackground="#97c7f1")
        self.back_button.place(x=-19, y=85)
        self.back_button.config(bg="#97c7f1")

        # This spacer allows for a gap between the 'Generate Member Report' banner and the 'First Name:' label.
        self.spacer = Label(self, text="", font=("", 45))
        self.spacer.pack()
        self.spacer.config(bg="#97c7f1")

        # This block creates the 'First Name:' label widget.
        self.first_name_label = Label(self, text="First Name:", font=("", 17))
        self.first_name_label.pack()
        self.first_name_label.config(bg="#97c7f1")

        # This segment creates the entry box to which the staff member inputs the first name of the member they
        # wish to generate the report on.
        self.first_name = Entry(self, bd=5, font=("", 17))
        self.first_name.pack()

        # This spacer creates a gap between the first name entry box and the last name label widget.
        self.spacer1 = Label(self, text="", font=("", 6))
        self.spacer1.pack()
        self.spacer1.config(bg="#97c7f1")

        # This block creates the 'Last Name:' label widget.
        self.last_name_label = Label(self, text="Last Name:", font=("", 17))
        self.last_name_label.pack()
        self.last_name_label.config(bg="#97c7f1")

        # This segment creates the entry box to which the staff member inputs the last name of the member they
        # wish to generate the report on.
        self.last_name = Entry(self, bd=5, font=("", 17))
        self.last_name.pack()

        # This spacer creates a gap between the last name entry box and the 'Generate' button.
        self.spacer2 = Label(self, text="", font=("", 12))
        self.spacer2.pack()
        self.spacer2.config(bg="#97c7f1")

        # This segment of code implements the 'Generate' button onto the window. When the button is pressed by the
        # staff associate, the 'member_report' method is executed.
        self.generate_photo = PhotoImage(file="Images/Generate Button.png")
        self.generate_button = Button(self, image=self.generate_photo, command=self.member_report,
                        height=60, width=200, borderwidth=0, activebackground="#97c7f1")
        self.generate_button.pack()
        self.generate_button.config(bg="#97c7f1")

        # This spacer creates a gap between the 'Generate' button and the 'answer' label widget.
        self.spacer2 = Label(self, text="", font=("", 12))
        self.spacer2.pack()
        self.spacer2.config(bg="#97c7f1")

        # The 'answer' label widget is created in order for the result of the member report request to be outputted to
        # the staff associate via the graphical interface.
        self.answer = Label(self, text="", font=("", 15))
        self.answer.pack()
        self.answer.config(bg="#97c7f1")

    # The 'member_report' method generates and saves the 'Member Report' documentation. The method is only executed when
    # the staff associate presses the 'Generate Member Report' button.
    def member_report(self):
        """METHOD FOR IMPLEMENTING ALL CUSTOMER INFORMATION"""
        os.chdir(r"C:\Users\Alex\PycharmProjects\Crook Log Leisure Centre\Member Reports")
        connection.connect(user="root", password="")
        first_name = self.first_name.get()
        last_name = self.last_name.get()
        document = Document()
        query = "SELECT * FROM users WHERE First_Name = %s AND Last_Name = %s"
        cursor.execute(query, (first_name, last_name,))
        result = cursor.fetchall()
        if result:
            for (Customer_ID, Title, First_Name, Last_Name, DOB, Email, Gender, Address, Post_Code, City, Membership_Type, Password, Deleted) \
                    in result:
                heading = ("%s %s - Member Report" % (first_name, last_name))
                paragraph = document.add_heading(heading.title())
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                document.add_paragraph \
                ("\nCustomer ID: {}\n\nTitle: {} \n\nFirst Name: {} \n\nLast Name: {} \n\nDate Of Birth: {} \n\nEmail: {} \n\nGender: {} \n\nAddress: {} \
                \n\nPost Code: {} \n\nCity: {} \n\nMembership Type: {} \n\nPassword: {} \n\nStatus: {}".format(Customer_ID, Title, First_Name, Last_Name, DOB,
                                                                              Email, Gender, Address, Post_Code, City, Membership_Type,
                                                                              Password, Deleted))
                name = ("%s %s - Report" % (first_name, last_name))
                file_type = (".docx")
                document.save(name.title()+file_type)
                self.answer.config(text="Completed.")
                connection.close()
                self.first_name.delete(0, 'end')
                self.last_name.delete(0, 'end')
                self.first_name.focus_set()
        else:
            self.answer.config(text="User Not Found in the Database.")

    # The 'executive_report' function is executed when the Administrator presses the 'Generate Executive Report' button
    # on the 'Administrator Menu' page.
    def executive_report():
        connection.connect(user="root", password="")
        os.chdir(r"C:\Users\Alex\PycharmProjects\Crook Log Leisure Centre\Executive Reports")
        document = Document()
        paragraph = document.add_heading("Executive Report")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        query = "SELECT COUNT(Role_Type_ID) from user_roles WHERE Role_Type_ID = 2;"
        cursor.execute(query)
        result = cursor.fetchall()
        document.add_paragraph(("\nNumber of Staff Associates: %s") % (result[0][0]))

        query = "SELECT COUNT(User_ID) from users WHERE Status = 'ACTIVE';"
        cursor.execute(query)
        result = cursor.fetchall()
        document.add_paragraph(("\nTotal Number of Active Members: %s\n") % (result[0][0]))
        ##

        query = "SELECT COUNT(Membership_Type) from users WHERE Membership_Type = 'Gym' AND Status = 'ACTIVE';"
        cursor.execute(query)
        result = cursor.fetchall()
        document.add_paragraph(("Gym Memberships: %s") % (result[0][0]))
        ##

        query = "SELECT COUNT(Membership_Type) from users WHERE Membership_Type = 'Swim23' AND Status = 'ACTIVE';"
        cursor.execute(query)
        result = cursor.fetchall()
        document.add_paragraph(("Swim23 Memberships: %s\n") % (result[0][0]))
        ##

        query = "SELECT COUNT(Membership_Type) from users WHERE Status = 'INACTIVE';"
        cursor.execute(query)
        result = cursor.fetchall()
        document.add_paragraph(("Total Number of Inactive Memberships: %s\n") % (result[0][0]))

        query = "SELECT COUNT(Gender) FROM users WHERE Gender = 'Male' AND Status = 'ACTIVE';"
        cursor.execute(query)
        result = cursor.fetchall()
        document.add_paragraph(("Male Members: %s") % (result[0][0]))

        query = "SELECT COUNT(Gender) FROM users WHERE Gender = 'Female' AND Status = 'ACTIVE';"
        cursor.execute(query)
        result = cursor.fetchall()
        document.add_paragraph(("Female Members: %s\n") % (result[0][0]))

        query = "SELECT MAX(DOB) FROM users WHERE Gender = 'Female' AND Status = 'ACTIVE';"
        cursor.execute(query)
        result = cursor.fetchall()
        document.add_paragraph(("Youngest Female Member: %s") % (result[0][0]))

        query = "SELECT MIN(DOB) FROM users WHERE Gender = 'Female' AND Status = 'ACTIVE';"
        cursor.execute(query)
        result = cursor.fetchall()
        document.add_paragraph(("Oldest Female Member: %s\n") % (result[0][0]))

        query = "SELECT MAX(DOB) FROM users WHERE Gender = 'Male' AND Status = 'ACTIVE';"
        cursor.execute(query)
        result = cursor.fetchall()
        document.add_paragraph(("Youngest Male Member: %s") % (result[0][0]))

        query = "SELECT MIN(DOB) FROM users WHERE Gender = 'Male' AND Status = 'ACTIVE';"
        cursor.execute(query)
        result = cursor.fetchall()
        document.add_paragraph(("Oldest Male Member: %s\n") % (result[0][0]))

        query = "SELECT COUNT(1) - COUNT(Exit_Time) FROM live_session WHERE Exit_Time is NULL;"
        cursor.execute(query)
        result = cursor.fetchall()
        document.add_paragraph(("Total Number of Members Currently Inside the Centre: %s") % (result[0][0]))

        connection.commit()
        name = ("Executive Report")
        file_type = (".docx")
        document.save(name.title() + file_type)
        connection.close()
        tkinter.messagebox.showinfo("Executive Report", "Executive Report Generated.")

    # The 'emergency_report' function is executed when the Administrator presses the 'Generate Emergency Report' button
    # on the 'Administrator Menu' page.
    def emergency_report():
        connection.connect(user="root", password="")
        os.chdir(r"C:\Users\Alex\PycharmProjects\Crook Log Leisure Centre\Executive Reports")
        document = Document()
        heading = ("Emergency Report\n")
        paragraph = document.add_heading(heading.title())
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name = ("Emergency Report")
        file_type = (".docx")
        query = ("SELECT User_ID FROM live_session WHERE Exit_Time is NULL;")
        cursor.execute(query)
        result = cursor.fetchall()
        for (User_ID) in result:
            query = "SELECT First_Name, Last_Name FROM users WHERE User_ID = %s"
            cursor.execute(query, (User_ID[0],))
            connection.commit()
            for (First_Name, Last_Name) in cursor:
                document.add_paragraph("First Name: {}\nLast Name: {}\n".format(First_Name, Last_Name))
        document.save(name.title()+file_type)
        tkinter.messagebox.showinfo("Emergency Report", "Emergency Report Generated.")

